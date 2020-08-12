'''Программа автоматизированного заполнения списка увольняемых.
Версия 1.0 (с графическим пользовательским интерфейсом).
'''

# Импорты из стандартной бибилиотеки.
import pickle
import random
import tkinter
import tkinter.filedialog
import tkinter.messagebox
# Импорты сторонних модулей.
import docx
import openpyxl


def get_operators_data_base():
    '''Функция возвращает базу данных операторов,
    находящуюся в том же каталоге,
    в котором запущена программа.
    '''
    # Файл, в котором находится база данных операторов.
    file_name = '.\Operators_data_base.data'
    # Загрузка базы данных операторов.
    with open(file_name, 'rb') as out:
        operators_data_base = pickle.load(out)
    return operators_data_base


def get_officers_data_base():
    '''Функция возвращает базу данных офицеров,
    находящуюся в том же каталоге,
    в котором запущена программа.
    '''
    # Файл, в котором находится база данных офицеров.
    file_name = '.\Officers_data_base.data'
    # Загрузка базы данных офицеров.
    with open(file_name, 'rb') as out:
        officers_data_base = pickle.load(out)
    return officers_data_base


def place_randomizer():
    '''Функция возвращает случайное место убытия увольняемого
    (из заданного в функции списка мест).
    '''
    places = [
        'ТЦ "Торговый центр"',
        'Кинотеатр "Домкино"',
        'Набережная'
        ]
    return random.choice(places)


def create_document():
    '''Функция формирования текстовой информации
    и заполнения документа.
    '''
    # Подгрузка базы данных операторов.
    operators_data_base = get_operators_data_base()
    # Сортировка списков увольняемых.
    general_operator_list.sort()
    post_operator_list.sort()
    church_operator_list.sort()
    # Вычисление общего числа строк основной таблицы увольняемых.
    total_rows = (
        len(general_operator_list)
        + len(post_operator_list) + 1
        )
    # Вычисление числа строк,
    # необходимых только для увольняемых в город.
    general_rows = len(general_operator_list) + 1
    # Импортирование шаблонного документа,
    # в котором заранее прописаны стили форматирования
    # (документ находится в том же каталоге, 
    # в котором запущена программа).
    doc = docx.Document('.\Template.docx')
    # Удаление первого пустого абзаца
    # (для того, чтобы начать заполнение документа с первой строчки).
    doc._body.clear_content()
    # Создание словаря для перевода численного значения месяца
    # в строковый формат.
    month_dictionary = {
        '01':'января',
        '02':'февраля',
        '03':'марта',
        '04':'апреля',
        '05':'мая',
        '06':'июня',
        '07':'июля',
        '08':'августа',
        '09':'сентября',
        '10':'октября',
        '11':'ноября',
        '12':'декабря',
        }
    # Формирование даты.
    date = (
        day_entry.get() + ' ' 
        + month_dictionary[month_entry.get()] + ' ' 
        + year_entry.get() + ' года'
        )
    # Добавление заголовка в создаваемый документ.
    doc.add_heading(
        'Cписок \n личного состава подразделения,\n'
        + 'убывающего в увольнение '
        + date, 
        1
        )
    # Формирование таблицы: 
    # число строк - по числу увольняемых в город и на почту,
    # число колонок - согласно утвержденному формату,
    # стиль - "Сетка таблицы" 
    # (в шаблонном документе заранее прописаны 
    # настройки форматирования содержимого таблицы).
    table = doc.add_table(rows=total_rows, cols=7, style='Table Grid')
    table.autofit = False
    # Создание списка 
    # с вычисленной шириной колонок таблицы в единицах EMU 
    # (вычислено по образцовому документу,
    # исходя из соотношения 1 дюйм = 914400 EMU, 1дюйм = 2,54 см).
    columns_size = [
        266090, 655168,
        1353312, 1349654,
        1440180, 896112, 905256
        ]
    # Задание ширины колонок через размер ячеек.
    for i in range(7):
        for j in range(total_rows):
            table.cell(j,i).width = columns_size[i]

    # Заполнение заголовочной части таблицы.
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'N п/п'
    hdr_cells[1].text = 'Воинское \n звание'
    hdr_cells[2].text = 'Фамилия \n инициалы'
    hdr_cells[3].text = 'Номер телефона'
    hdr_cells[4].text = 'Место убытия'
    hdr_cells[5].text = 'До которого \n часа и дня \n уволен'
    hdr_cells[6].text = 'Время \n прибытия'

    # Заполнение общего списка увольняемых.
    for i in range(1, general_rows):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)+'.'
        row_cells[1].text = operators_data_base[general_operator_list[i-1]][0]
        row_cells[2].text = general_operator_list[i-1]
        row_cells[3].text = operators_data_base[general_operator_list[i-1]][1]
        row_cells[4].text = place_randomizer()
        row_cells[5].text = '20:00'
    # Заполнение списка увольняемых на почту.
    if post_operator_list != []:
        # k - индекс для обращения к элементам списка увольняемых на почту.
        k = 0
        for i in range(general_rows, total_rows):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(i)+'.'
            row_cells[1].text = operators_data_base[post_operator_list[k]][0]
            row_cells[2].text = post_operator_list[k]
            row_cells[3].text = operators_data_base[post_operator_list[k]][1]
            row_cells[4].text = 'Отделение почты'
            row_cells[5].text = '16:00'
            k += 1
    
    # Заполнение списка увольняемых в Храм.
    if church_operator_list != []:
        # Вычисление числа строк для таблицы увольняемых в Храм.
        church_rows = len(church_operator_list) + 1
        # Добавление заголовка таблицы.
        doc.add_heading(
            'Cписок \n личного состава подразделения,\n' 
            + 'убывающего на службу в Храм \n'
            + date,
            1
            )
        # Формирование таблицы - аналогично основной таблице увольняемых.
        table1 = doc.add_table(rows=church_rows, cols=7, style='Table Grid')
        table1.autofit = False
        columns_size1 = [
            266090, 655168,
            1353312,1349654, 
            1440180,896112,905256
            ]
        for i in range(7):
            for j in range(church_rows):
                table1.cell(j,i).width = columns_size1 [i]
        # Заполнение заголовочной части таблицы.
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = 'N п/п'
        hdr_cells[1].text = 'Воинское \n звание'
        hdr_cells[2].text = 'Фамилия \n инициалы'
        hdr_cells[3].text = 'Номер телефона'
        hdr_cells[4].text = 'С которого \n часа \n уволен'
        hdr_cells[5].text = 'До которого \n часа и дня \n уволен'
        hdr_cells[6].text = 'Время \n прибытия'
        # k - индекс для общей нумерации количества увольняемых
        # при переходе от основной таблицы в таблице увольняемых в Храм.
        k = total_rows
        for i in range(1, church_rows):
            row_cells = table1.rows[i].cells
            row_cells[0].text = str(k)+'.'
            row_cells[1].text = operators_data_base[church_operator_list[i-1]][0]
            row_cells[2].text = church_operator_list[i-1]
            row_cells[3].text = operators_data_base[church_operator_list[i-1]][1]
            row_cells[4].text = '7:00'
            row_cells[5].text = '12:00'
            k += 1

    # Заполнение таблицы для подписи командира роты.
    doc.add_paragraph()
    doc.add_paragraph()
    # Формирование таблицы со скрытыми границами:
    # стиль "Светлая сетка" 
    # (в шаблонном документе заранее прописаны
    # настройки форматирования содержимого таблицы).
    table2 = doc.add_table(rows=4, cols=2, style='Light Grid')
    table2.autofit = False
    # Формирование даты подписи документа командиром роты.
    date_sign = (
        sign_day_entry.get() + ' ' 
        + month_dictionary[sign_month_entry.get()] + ' ' 
        + sign_year_entry.get() + ' года'
        )
    # Задание ширины колонок таблицы.
    columns_size2 = [3328416, 3511296]
    for i in range(2):
        for j in range(3):
            table2.cell(j,i).width = columns_size2 [i]
    # Подгрузка базы данных офицеров.
    officers_data_base = get_officers_data_base()
    # Заполнение таблицы.
    table2.rows[0].cells[1].text = officers_data_base[army_signer_flag.get()][1]
    table2.rows[1].cells[1].text = officers_data_base[army_signer_flag.get()][2]
    table2.rows[3].cells[0].text = date_sign
    table2.rows[3].cells[1].text = (
        officers_data_base[army_signer_flag.get()][0]
        + 39*' ' + army_signer_flag.get()
        )

    # Заполнение таблицы для подписи ЗНТ по ВПР.
    doc.add_paragraph()
    doc.add_paragraph()
    # Формирование таблицы со скрытыми границами
    # стиль "Светлая сетка" 
    # (в шаблонном документе заранее прописаны
    #  настройки форматирования содержимого таблицы).
    table3 = doc.add_table(rows=4, cols=2, style='Light Grid')
    table3.autofit = False
    # Задание ширины колонок таблицы.
    columns_size3 = [4230014, 2609698]
    for i in range(2):
        for j in range(3):
            table3.cell(j,i).width = columns_size3[i]
    # Заполнение таблицы.
    table3.rows[0].cells[0].text = officers_data_base[policy_signer_flag.get()][1]
    table3.rows[3].cells[0].text = (
        officers_data_base[policy_signer_flag.get()][0] 
        + 54*' ' + policy_signer_flag.get()
        )

    # Формирование строки названия документа.
    document_name = 'Увольняемые '+ date +'.docx'
    # Выбор папки, в которую будет сохранен документ.
    folder_name = tkinter.filedialog.askdirectory(
        title='Выберите папку для сохранения файла'
        )
    # Защита от пустого имени папки.
    if folder_name == '':
        return
    # Сохранение документа.
    doc.save(folder_name +'\\' + document_name)
    tkinter.messagebox.showinfo('Завершение процесса','Файл ' + document_name + ' успешно сохранен')
    # Очищение буферов хранения списков увольняемых,
    # для возможного дальнейшего использования программы.
    general_operator_list.clear()
    general_operator_list_indicator.config(
        text='Основной список увольняемых не заполнен!',
        bg='#ff9d9d',
        )
    church_operator_list.clear()
    church_operator_list_indicator.config(
        text='Cписок увольняемых \nв Храм \nне заполнен!',
        bg='#ff9d9d',
        )
    post_operator_list.clear()
    post_operator_list_indicator.config(
        text='Cписок увольняемых \nна почту \nне заполнен!',
        bg='#ff9d9d',
        )
   

# Графический интерфейс пользователя.
# Функции, описывающие события,
# происходящие по клику на кнопки интерфейса.

def show_operators_listboxes(title):
    '''Функция открывает окно со
    списком всех операторов 2 научной роты.
    Запускается нажатием кнопок следующих кнопок:
    - fill_general_list_by_hand,
    - fill_church_list,
    - fill_post_list.
    В функцию в качестве параметра передаётся название окна.
    '''  

    def add_in_list():
        '''Функция добавляет выбранных операторов
        в определённый буфер по нажатию кнопки ok_button. 
        '''
        # Выбор буфера 
        # в зависимости от переданного 
        # в функцию show_operators_listboxes
        # названия окна.
        # Буферы предварительно очищаются 
        # (на случай, если в них есть содержимое 
        # после неудачной попытки заполнения).
        for operator in check_operators_listbox.get(0, tkinter.END):
            if title == 'Заполнение основого списка':
                general_operator_list.append(operator)
                general_operator_list_indicator.config(
                    text='Основной список увольняемых заполнен!',
                    bg='#7dff7d',    
                    )
            elif title == 'Заполнение cписка в Храм':
                church_operator_list.append(operator)
                church_operator_list_indicator.config(
                    text='Cписок увольняемых \nв Храм \nзаполнен!',
                    bg='#7dff7d',    
                    )
            elif title == 'Заполнение списка на почту':
                post_operator_list.append(operator)
                post_operator_list_indicator.config(
                    text='Cписок увольняемых \nна почту \nзаполнен!',
                    bg='#7dff7d',    
                    )
        # Закрытие окна со списком операторов.
        operators_listbox_window.destroy()

    def clear_check_operators_listbox():
        '''Функция очищает список выбранных операторов
        по нажатию кнопки clear_button
        '''
        check_operators_listbox.delete(0, tkinter.END)

    def add_in_check_operators_listbox():
        '''Функция добавляет выбранного оператора
        в виджет-список check_operators_listbox
        по нажатию кнопки add_button
        '''
        # Определение индексов выделенных операторов.
        select_operators_index = list(operators_listbox.curselection())
        # Добавление операторов в виджет-список check_operators_listbox.
        for i in select_operators_index:
           check_operators_listbox.insert(
               tkinter.END, 
               operators_listbox.get(i)
               ) 

    # Очищение буферов хранения списков увольняемых.
    if title == 'Заполнение основого списка':
        general_operator_list.clear()
        general_operator_list_indicator.config(
            text='Основной список увольняемых не заполнен!',
            bg='#ff9d9d',
            )
    elif title == 'Заполнение cписка в Храм':
        church_operator_list.clear()
        church_operator_list_indicator.config(
            text='Cписок увольняемых \nв Храм \nне заполнен!',
            bg='#ff9d9d',
            )
    elif title == 'Заполнение списка на почту':
            post_operator_list.clear()
            post_operator_list_indicator.config(
            text='Cписок увольняемых \nна почту \nне заполнен!',
            bg='#ff9d9d',
            )
    # Создание окна со списком операторов и его виджетов.
    operators_listbox_window = tkinter.Toplevel()
    operators_listbox_window.title(title)
    operators_listbox_window.config(bg='#cfcfcf')
    data_base_label = tkinter.Label(
        operators_listbox_window, 
        text='База данных:', 
        bg='#cfcfcf',
        )
    data_base_label.grid(row=0, column=0)
    check_listbox_label = tkinter.Label(
        operators_listbox_window, 
        text='Выбрано:', 
        bg='#cfcfcf',
        )
    check_listbox_label.grid(row=0, column=3)
    # Создание и разщмещение виджета-списка 
    # операторов из базы данных.
    operators_listbox = tkinter.Listbox(
        operators_listbox_window,
        selectmode=tkinter.EXTENDED,
        )
    operators_listbox.grid(row=1, column=0)
    # Добавление и размещение скроллера.
    scroll_1 = tkinter.Scrollbar(
        operators_listbox_window,
        command=operators_listbox.yview,
        )
    scroll_1.grid(
        row=1, 
        column=1, 
        sticky=tkinter.W + tkinter.N + tkinter.S,
        )
    # Установка скроллера виджету-списку.
    operators_listbox.config(yscrollcommand=scroll_1.set)
    # Получение пути к файлу, 
    # в котором находится база данных операторов.
    file_name = '.\Operators_data_base.data'
    # Подгрузка базы данных операторов.
    with open(file_name, 'rb') as out:
        data_base = pickle.load(out)
    # Добавление операторов из базы данных
    # в виджет-список.
    for i in data_base:
        operators_listbox.insert(tkinter.END, i)
    # Создание и размещение кнопки
    # для добавления выбранных операторов
    # в виджет-список выбранных операторов.
    add_button = tkinter.Button(
        operators_listbox_window,
        text='Добавить', 
        command=add_in_check_operators_listbox,
        )
    add_button.grid(row=1, column=2)
    # Создание и разщмещение виджета-списка 
    # выбранных операторов.
    check_operators_listbox = tkinter.Listbox(operators_listbox_window)
    check_operators_listbox.grid(row=1, column=3)
    # Добавление и размещение скроллера.
    scroll_2 = tkinter.Scrollbar(
        operators_listbox_window,
        command=check_operators_listbox.yview,
        )
    scroll_2.grid(
        row=1, 
        column=4, 
        sticky=tkinter.W + tkinter.N + tkinter.S,
        )
    # Установка скроллера виджету-списку.
    check_operators_listbox.config(yscrollcommand=scroll_2.set)
    # Создание и размещение кнопки 
    # добавления операторов из виджета-списка 
    # выбранных операторов
    # в соответствующий буфер. 
    ok_button = tkinter.Button(
        operators_listbox_window,
        text='Принять',
        bg='#ff9d3c', 
        command=add_in_list,
        ) 
    ok_button.grid(row=2, column=1, columnspan=2)
    clear_button = tkinter.Button(
        operators_listbox_window,
        text='Очистить', 
        command=clear_check_operators_listbox,
        )
    clear_button.grid(row=2, column=3) 


def show_add_file_window():
    '''Функция открывает меню для прикрепления
    файла с графиком увольняемых 
    по нажатию кнопки fill_general_list_from_file.
    '''

    def add_file():
        '''Функция определяет путь 
        к файлу с графиком увольняемых
        по нажатию кнопки add_file_button.
        '''
        add_file_entry.delete(0, tkinter.END)
        add_file_entry.insert(0, 
            tkinter.filedialog.askopenfilename()
            )

    def parse_operators():
        '''Функция заполняет основной список увольняемых,
        получая данные из прикреплённого файла
        по нажатию кнопки parse_operators_button.
        '''
        # Создание словаря перевода численного значения месяца
        # в формат, принятый в файле с графиком увольняемых.
        month_dictionary = {
            1:'ЯНВАРЬ',
            2:'ФЕВРАЛЬ',
            3:'МАРТ',
            4:'АПРЕЛЬ',
            5:'МАЙ',
            6:'ИЮНЬ',
            7:'ИЮЛЬ',
            8:'АВГУСТ',
            9:'СЕНТЯБРЬ',
            10:'ОКТЯБРЬ',
            11:'НОЯБРЬ',
            12:'ДЕКАБРЬ',
            13: 'ЯНВАРЬ',
            }
        # Открытие excel-файла с графиком увольняемых. 
        workbook = openpyxl.load_workbook(
            add_file_entry.get(),
            data_only=True
            )
        # Выбор листа.
        worksheet = workbook['План']
        # Получение месяца и даты увольнения.
        month = month_dictionary[int(month_entry.get())]
        day = int(day_entry.get())
        # Создание флага для организации процесса поиска.
        flag = True
        # Установление начального значения счётчика колонок
        # (согласно формату файла с графиком,
        # данные начинаются с 4 колонки)
        columns_counter = 4
        # Определение строк для поиска месяца и дня
        # (согласно формату файла с графиком).
        month_row = 2
        day_row = 4
        # Поиск столбца, соответствующего искомому месяцу.
        # (процесс поиска идёт до тех пор,
        # пока в 45 строке не появится пустое значение,
        # что свидетельствует о выходе за пределы 
        # области интересующих данных).
        while flag:
            value_month = worksheet.cell(
                row=month_row, 
                column=columns_counter
                ).value
            if value_month == month:
                break
            elif worksheet.cell(row=45, column=columns_counter).value == None:
                flag = False
                tkinter.messagebox.showerror(
                    'Ошибка',
                    'Указанный месяц в файле не найден'
                    )
                break
            columns_counter += 1
        # Поиск столбца, соответствующего искомому дню.
        while flag:
            value_day = worksheet.cell(
                row=day_row,
                column=columns_counter
                ).value
            if value_day == day:
                break
            # Проверка на наличие даты увольнения в файле с графиком,
            # ошибка возникает если:
            # - достигнута граница области данных
            # (в строке 5 достигнут 
            # первый столбец с итоговой суммой увольнений);
            # - достигнута граница введённого месяца;
            # - считанная дата превосходит искомую.
            elif (
                    type(
                        worksheet.cell(
                            row=5, 
                            column=columns_counter + 1
                            ).value
                        ) == int
                    or worksheet.cell(
                        row=month_row,
                        column=columns_counter + 1
                        ).value 
                        == month_dictionary[int(month_entry.get())+1]
                    or value_day > day
                    ):
                tkinter.messagebox.showerror(
                    'Ошибка',
                    'Указанный день в файле не найден'
                    )
                flag = False
                break
            columns_counter += 1
        # Поиск увольняемых по графику.
        # Установление начального значения счётчика строк
        # (согласно формату файла с графиком,
        # фамилии операторов начинаются с 5 строки).
        rows_counter = 5
        while flag and worksheet.cell(row=rows_counter, column=3).value != None:
            cell = worksheet.cell(
                row= rows_counter, 
                column=columns_counter
                ).value
            if cell == 'У':
                general_operator_list.append(
                    worksheet.cell(row= rows_counter, column=3).value
                    )
                general_operator_list_indicator.config(
                    text='Основной список увольняемых заполнен!',
                    bg='#7dff7d',    
                    )
            rows_counter += 1
        # Закрытие окна прикрепления файла с графиком увольняемых.
        add_file_window.destroy()
    
    # Очищение буфера хранения основного списка увольняемых.
    general_operator_list.clear()
    general_operator_list_indicator.config(
        text='Основной список увольняемых не заполнен!',
        bg='#ff9d9d',
        )
    # Создание и размещение 
    # окна прикрепления файла с графиком увольняемых.
    add_file_window = tkinter.Toplevel()
    add_file_window.title('Прикрепление файла с графиком')
    add_file_window.config(bg='#cfcfcf')
    add_file_entry = tkinter.Entry(add_file_window, width=100)
    add_file_entry.grid(row=0, column=0)
    add_file_button = tkinter.Button(
        add_file_window, 
        text='Указать путь', 
        command=add_file,
        )
    add_file_button.grid(row=0, column=1)
    parse_operators_button = tkinter.Button(
        add_file_window, 
        text='Заполнить список', 
        bg='#ff9d3c', 
        command=parse_operators,
        )
    parse_operators_button.grid(row=1, column=0, columnspan=2)


# Буферы для хранения списков увольняемых.
general_operator_list = []
church_operator_list = []
post_operator_list = []
# Создание главного окна и его виджетов.
root = tkinter.Tk()
root.title('Увольнения 1.0')
root.config(bg='#cfcfcf')
frame_1 = tkinter.LabelFrame(
    text='Заполнение основного списка увольняемых', 
    bg='#cfcfcf',
    )
frame_1.grid(row=0, column=0, padx=5)
day_label = tkinter.Label(frame_1, text='Дата', bg='#cfcfcf')
day_label.grid(row=0, column=1)
day_entry = tkinter.Entry(frame_1, width=5, justify=tkinter.CENTER)
day_entry.grid(row=1, column=1)
month_label = tkinter.Label(frame_1, text='Месяц', bg='#cfcfcf')
month_label.grid(row=0, column=2)
month_entry = tkinter.Entry(frame_1, width=5, justify=tkinter.CENTER)
month_entry.grid(row=1, column=2)
year_label = tkinter.Label(frame_1, text='Год', bg='#cfcfcf')
year_label.grid(row=0, column=3)
year_entry = tkinter.Entry(frame_1, width=5, justify=tkinter.CENTER)
year_entry.grid(row=1, column=3)
# Создание и рамещение кнокпки заполнения 
# основного списка увольняемых вручную
# (в функцию действия по на нажатию
# с помощью lamda-функции пробрасывается 
# название открывающегося окна).
fill_general_list_by_hand = tkinter.Button(
    frame_1,
    text='Заполнить список \n вручную',
    bg='#ff9d3c',
    command=(
        lambda title='Заполнение основого списка': 
        show_operators_listboxes(title)
        ),
    )
fill_general_list_by_hand.grid(
    row=2, 
    column=0, 
    columnspan=2, 
    pady=5, 
    padx=5,
    )
fill_general_list_from_file = tkinter.Button(
    frame_1,
    text='Заполнить список \n по графику',
    bg='#ff9d3c',
    command=show_add_file_window,
    )
fill_general_list_from_file.grid(row=2, column=3, columnspan=2, padx=5)
general_operator_list_indicator = tkinter.Label(frame_1,
    text='Основной список увольняемых не заполнен!',
    bg='#ff9d9d',
    )
general_operator_list_indicator.grid(
    row=3, 
    column=0, 
    columnspan=6, 
    pady=5,
    )
frame_2 = tkinter.LabelFrame(
    text='Заполнение опциональных списков увольняемых', 
    bg='#cfcfcf',
    )
frame_2.grid(row=0, column=1, padx=5, sticky=tkinter.N)
fill_church_list = tkinter.Button(
    frame_2,
    text='Заполнить список \n увольняемых в Храм',
    bg='#ff9d3c',
    command=(
        lambda title='Заполнение cписка в Храм': 
        show_operators_listboxes(title)
        ),
    )
fill_church_list.grid(row=0, column=0, padx=5)
church_operator_list_indicator = tkinter.Label(
    frame_2,
    text='Cписок увольняемых \nв Храм \nне заполнен!',
    bg='#ff9d9d',
    )
church_operator_list_indicator.grid(row=1, column=0)
fill_post_list = tkinter.Button(
    frame_2,
    text='Заполнить список \n увольняемых на почту',
    bg='#ff9d3c',
    command=(
        lambda title='Заполнение списка на почту': 
        show_operators_listboxes(title)
        ),
    )
fill_post_list.grid(row=0, column=1, pady=5)
post_operator_list_indicator = tkinter.Label(
    frame_2,
    text='Cписок увольняемых \nна почту \nне заполнен!',
    bg='#ff9d9d',
    )
post_operator_list_indicator.grid(row=1, column=1)
frame_3 = tkinter.LabelFrame(
    text='Заполнение служебной информации', 
    bg='#cfcfcf',
    )
frame_3.grid(row=1, column=0, columnspan=2)
sign_label = tkinter.Label(
    frame_3, 
    text='Дата подписи командиром роты:', 
    bg='#cfcfcf',
    )
sign_label.grid(row=0, column=0, columnspan=3)
sign_day_label = tkinter.Label(frame_3, text='Дата', bg='#cfcfcf')
sign_day_label.grid(row=1, column=0)
sign_day_entry = tkinter.Entry(frame_3, width=5, justify=tkinter.CENTER)
sign_day_entry.grid(row=2, column=0)
sign_month_label = tkinter.Label(frame_3, text='Месяц', bg='#cfcfcf')
sign_month_label.grid(row=1, column=1)
sign_month_entry = tkinter.Entry(frame_3, width=5, justify=tkinter.CENTER)
sign_month_entry.grid(row=2, column=1)
sign_year_label = tkinter.Label(frame_3, text='Год', bg='#cfcfcf')
sign_year_label.grid(row=1, column=2)
sign_year_entry = tkinter.Entry(frame_3, width=5, justify=tkinter.CENTER)
sign_year_entry.grid(row=2, column=2)
army_signer_label = tkinter.Label(
    frame_3, 
    text='Выберите подписывающего список в роте:', 
    bg='#cfcfcf',
    )
army_signer_label.grid(row=3, column=0, columnspan=4)
army_signer_flag = tkinter.StringVar()
army_signer_flag.set('М.Майоров')
army_signer_radiobutton_1 = tkinter.Radiobutton(
    frame_3,
    text='майор М.Майоров',
    variable=army_signer_flag,
    value='М.Майоров', 
    bg='#cfcfcf',
    )
army_signer_radiobutton_1.grid(
    row=4, 
    column=0, 
    columnspan=4, 
    sticky=tkinter.W,
    )
army_signer_radiobutton_2 = tkinter.Radiobutton(
    frame_3,
    text='капитан К.Капитанов',
    variable=army_signer_flag,
    value='К.Капитанов',
    bg='#cfcfcf',
    )
army_signer_radiobutton_2.grid(
    row=5, 
    column=0, 
    columnspan=4, 
    sticky=tkinter.W,
    )
army_signer_radiobutton_3 = tkinter.Radiobutton(
    frame_3,
    text='лейтенант Л.Лейтенантов',
    variable=army_signer_flag,
    value='Л.Лейтенантов',
    bg='#cfcfcf',
    )
army_signer_radiobutton_3.grid(
    row=6, 
    column=0, 
    columnspan=4, 
    sticky=tkinter.W,
    )
policy_signer_label = tkinter.Label(
    frame_3, 
    text='Выберите подписывающего список в отделе ВПР:', 
    bg='#cfcfcf',
    )
policy_signer_label.grid(row=0, column=5, columnspan=4)
policy_signer_flag = tkinter.StringVar()
policy_signer_flag.set('П.Полковников')
policy_signer_radiobutton_1 = tkinter.Radiobutton(frame_3,
    text='полковник П.Полковников',
    variable=policy_signer_flag,
    value='П.Полковников',
    bg='#cfcfcf',
    )
policy_signer_radiobutton_1.grid(
    row=1, 
    column=5, 
    columnspan=4, 
    sticky=tkinter.W,
    )
policy_signer_radiobutton_2 = tkinter.Radiobutton(
    frame_3,
    text='майор С.Сидоров',
    variable=policy_signer_flag,
    value='С.Сидоров',
    bg='#cfcfcf',
    )
policy_signer_radiobutton_2.grid(
    row=2, 
    column=5, 
    columnspan=4, 
    sticky=tkinter.W,
    )
policy_signer_radiobutton_3 = tkinter.Radiobutton(
    frame_3,
    text='майор С.Сергеев',
    variable=policy_signer_flag,
    value='С.Сергеев',
    bg='#cfcfcf',
    )
policy_signer_radiobutton_3.grid(
    row=3, 
    column=5, 
    columnspan=4, 
    sticky=tkinter.W,
    )
policy_signer_radiobutton_4 = tkinter.Radiobutton(
    frame_3,
    text='Дежурный',
    variable=policy_signer_flag,
    value='Дежурный',
    bg='#cfcfcf',
    )
policy_signer_radiobutton_4.grid(
    row=4, 
    column=5, 
    columnspan=4, 
    sticky=tkinter.W)
generate_file = tkinter.Button(
    text="Сгенерировать файл",
    bg='#ff9d3c',
    command=create_document,
    )
generate_file.grid(row=2, column=0, columnspan=2)
root.mainloop()